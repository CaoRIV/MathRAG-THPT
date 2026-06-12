from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import select

from app.core.config import get_settings
from app.db.models import Document, Exam, User
from app.db.session import SessionLocal
from app.main import app


def auth_headers(client: TestClient) -> dict[str, str]:
    settings = get_settings()
    response = client.post(
        "/api/v1/auth/login",
        json={
            "email": settings.admin_email,
            "password": settings.admin_password,
        },
    )
    assert response.status_code == 200
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


def test_health_and_chat_contracts() -> None:
    with TestClient(app) as client:
        headers = auth_headers(client)
        health = client.get("/api/v1/health")
        assert health.status_code == 200
        response = client.post(
            "/api/v1/chat",
            json={"message": "Công thức nguyên hàm của x^n là gì?", "mode": "study"},
            headers=headers,
        )
        assert response.status_code == 200
        payload = response.json()
        assert payload["answer"]
        assert payload["citations"]
        assert payload["citations"][0]["chunk_id"]


def test_exam_mode_defaults_to_progressive_hint() -> None:
    with TestClient(app) as client:
        headers = auth_headers(client)
        response = client.post(
            "/api/v1/chat",
            json={"message": "Tìm khoảng đồng biến của hàm số", "mode": "exam"},
            headers=headers,
        )
        assert response.status_code == 200
        assert response.json()["assistance_level"] == "hint"


def test_search_document_topics_and_quiz() -> None:
    with TestClient(app) as client:
        headers = auth_headers(client)
        search = client.post(
            "/api/v1/search",
            json={"query": "logarit đổi cơ số", "include_scores": True},
            headers=headers,
        )
        assert search.status_code == 200
        assert search.json()["items"]
        document_id = search.json()["items"][0]["document_id"]
        assert client.get(
            f"/api/v1/documents/{document_id}",
            headers=headers,
        ).status_code == 200
        assert client.get("/api/v1/topics", headers=headers).status_code == 200
        quiz = client.post(
            "/api/v1/quizzes/generate",
            json={"topic": "Hàm số", "question_count": 1},
            headers=headers,
        )
        assert quiz.status_code == 200
        assert len(quiz.json()["questions"]) == 1


def test_registration_login_and_role_protection() -> None:
    email = f"student-{uuid4()}@example.com"
    with TestClient(app) as client:
        register = client.post(
            "/api/v1/auth/register",
            json={
                "email": email,
                "full_name": "Học sinh tích hợp",
                "password": "StrongPass123!",
            },
        )
        assert register.status_code == 201
        payload = register.json()
        assert payload["user"]["role"] == "user"
        user_headers = {"Authorization": f"Bearer {payload['access_token']}"}
        assert client.get("/api/v1/auth/me", headers=user_headers).status_code == 200
        forbidden = client.get("/api/v1/admin/documents", headers=user_headers)
        assert forbidden.status_code == 403
        assert client.post(
            "/api/v1/chat",
            json={"message": "Hàm số đồng biến khi nào?", "mode": "study"},
        ).status_code == 401
    with SessionLocal() as session:
        user = session.scalar(select(User).where(User.email == email))
        if user:
            session.delete(user)
            session.commit()


def test_admin_can_upload_docx_and_search_it() -> None:
    title = f"Chuyên đề cấp số cộng {uuid4()}"
    document = DocxDocument()
    document.add_heading("Công thức cấp số cộng", level=1)
    document.add_paragraph(
        "Số hạng tổng quát của cấp số cộng là $u_n=u_1+(n-1)d$."
    )
    content = BytesIO()
    document.save(content)

    document_id = None
    try:
        with TestClient(app) as client:
            headers = auth_headers(client)
            upload = client.post(
                "/api/v1/admin/documents/upload",
                headers=headers,
                files={
                    "file": (
                        "cap-so-cong.docx",
                        content.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                data={
                    "title": title,
                    "topic": "Dãy số",
                    "grade": "11",
                    "content_type": "formula",
                    "description": "Tài liệu kiểm thử upload.",
                },
            )
            assert upload.status_code == 201, upload.text
            document_id = upload.json()["id"]
            assert upload.json()["chunk_count"] >= 1
            listing = client.get("/api/v1/admin/documents", headers=headers)
            assert listing.status_code == 200
            assert any(item["title"] == title for item in listing.json()["items"])
            search = client.post(
                "/api/v1/search",
                headers=headers,
                json={
                    "query": title,
                    "filters": {"grade": 11},
                    "include_scores": True,
                },
            )
            assert search.status_code == 200
            assert any(item["title"] == title for item in search.json()["items"])
    finally:
        if document_id:
            with SessionLocal() as session:
                stored_document = session.get(Document, document_id)
                if stored_document:
                    source_path = Path(stored_document.source_path or "")
                    session.delete(stored_document)
                    session.commit()
                    source_path.unlink(missing_ok=True)


def test_admin_can_normalize_and_review_exam_questions() -> None:
    exam_id = None
    with TestClient(app) as client:
        headers = auth_headers(client)
        created = client.post(
            "/api/v1/admin/exams",
            headers=headers,
            json={
                "title": f"De thi thu THPT {uuid4()}",
                "year": 2026,
                "school": "Truong THPT Tich hop",
                "province": "Ha Noi",
                "exam_type": "mock",
                "duration_minutes": 90,
                "expected_question_count": 50,
                "grade": 12,
            },
        )
        assert created.status_code == 201, created.text
        exam_id = created.json()["id"]
        assert created.json()["question_count"] == 0
        assert created.json()["processing_status"] == "uploaded"

        question_payload = {
            "question_number": 1,
            "question_type": "multiple_choice",
            "prompt_markdown": "Tinh $\\int_0^1 x^2\\,dx$.",
            "options": [
                {"key": "A", "content_markdown": "$1/3$"},
                {"key": "B", "content_markdown": "$1/2$"},
                {"key": "C", "content_markdown": "$1$"},
                {"key": "D", "content_markdown": "$2$"},
            ],
            "correct_answer": "A",
            "solution_markdown": "$\\int_0^1 x^2\\,dx=1/3$.",
            "difficulty": "easy",
            "topics": ["Nguyen ham - Tich phan"],
            "formulas": [
                {
                    "raw_text": "integral from 0 to 1 of x squared",
                    "latex": "\\int_0^1 x^2\\,dx",
                    "normalized": "\\int_{0}^{1}x^{2}dx",
                }
            ],
            "page_number": 1,
            "extraction_status": "needs_review",
            "extraction_confidence": 0.96,
        }
        question = client.post(
            f"/api/v1/admin/exams/{exam_id}/questions",
            headers=headers,
            json=question_payload,
        )
        assert question.status_code == 201, question.text
        question_id = question.json()["id"]
        assert question.json()["options"][0]["key"] == "A"
        assert question.json()["formulas"][0]["normalized"]

        duplicate = client.post(
            f"/api/v1/admin/exams/{exam_id}/questions",
            headers=headers,
            json=question_payload,
        )
        assert duplicate.status_code == 422

        premature_approval = client.patch(
            f"/api/v1/admin/exams/{exam_id}",
            headers=headers,
            json={"processing_status": "approved"},
        )
        assert premature_approval.status_code == 422

        verified = client.patch(
            f"/api/v1/admin/exams/{exam_id}/questions/{question_id}",
            headers=headers,
            json={"extraction_status": "verified"},
        )
        assert verified.status_code == 200, verified.text

        approved = client.patch(
            f"/api/v1/admin/exams/{exam_id}",
            headers=headers,
            json={"processing_status": "approved"},
        )
        assert approved.status_code == 200, approved.text
        assert approved.json()["question_count"] == 1
        assert approved.json()["questions"][0]["question_number"] == 1

        listing = client.get("/api/v1/admin/exams", headers=headers)
        assert listing.status_code == 200
        assert any(item["id"] == exam_id for item in listing.json()["items"])

    if exam_id:
        with SessionLocal() as session:
            exam = session.get(Exam, exam_id)
            if exam:
                session.delete(exam)
                session.commit()


def test_exam_upload_auto_parses_and_reparse_preserves_verified_questions() -> None:
    title = f"De thi tu dong {uuid4()}"
    source = DocxDocument()
    source.add_heading("De thi thu tot nghiep THPT", level=1)
    source.add_paragraph(r"Câu 1. Tính $\int_0^1 x^2\,dx$.")
    source.add_paragraph(
        r"A. $\frac{1}{3}$ B. $\frac{1}{2}$ C. $1$ D. $2$"
    )
    source.add_paragraph("Câu 2. Hàm số nào sau đây đồng biến trên R?")
    source.add_paragraph("A. y=x^3 B. y=-x C. y=-x^2 D. y=1/x")
    source.add_heading("ĐÁP ÁN", level=1)
    source.add_paragraph("1 A 2 A")
    source.add_heading("LỜI GIẢI", level=1)
    source.add_paragraph(r"Câu 1. Ta có $\int_0^1 x^2\,dx=\frac{1}{3}$.")
    source.add_paragraph("Câu 2. Hàm số y=x^3 có đạo hàm không âm.")
    content = BytesIO()
    source.save(content)

    document_id = None
    exam_id = None
    try:
        with TestClient(app) as client:
            headers = auth_headers(client)
            upload = client.post(
                "/api/v1/admin/documents/upload",
                headers=headers,
                files={
                    "file": (
                        "de-thi.docx",
                        content.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                data={
                    "title": title,
                    "topic": "Đề thi THPT",
                    "grade": "12",
                    "content_type": "exam",
                },
            )
            assert upload.status_code == 201, upload.text
            payload = upload.json()
            document_id = payload["id"]
            exam_id = payload["exam_id"]
            assert payload["exam_processing_status"] == "needs_review", payload
            assert payload["exam_parse_report"]["detected_questions"] == 2
            assert payload["exam_parse_report"]["answers_matched"] == 2

            exam = client.get(
                f"/api/v1/admin/exams/{exam_id}",
                headers=headers,
            )
            assert exam.status_code == 200
            assert exam.json()["question_count"] == 2
            first_question = exam.json()["questions"][0]
            assert first_question["page_number"] is None
            assert first_question["correct_answer"] == "A"

            verified = client.patch(
                f"/api/v1/admin/exams/{exam_id}/questions/{first_question['id']}",
                headers=headers,
                json={
                    "extraction_status": "verified",
                    "prompt_markdown": "Nội dung đã được Admin kiểm duyệt.",
                },
            )
            assert verified.status_code == 200

            reparse = client.post(
                f"/api/v1/admin/documents/{document_id}/parse-exam",
                headers=headers,
            )
            assert reparse.status_code == 200, reparse.text
            assert reparse.json()["preserved_verified_questions"] == 1
            assert reparse.json()["updated_questions"] == 1

            after = client.get(
                f"/api/v1/admin/exams/{exam_id}",
                headers=headers,
            )
            assert after.status_code == 200
            assert after.json()["questions"][0]["prompt_markdown"] == (
                "Nội dung đã được Admin kiểm duyệt."
            )
    finally:
        with SessionLocal() as session:
            if exam_id:
                exam = session.get(Exam, exam_id)
                if exam:
                    session.delete(exam)
                    session.flush()
            if document_id:
                document = session.get(Document, document_id)
                if document:
                    source_path = Path(document.source_path or "")
                    session.delete(document)
                    session.commit()
                    source_path.unlink(missing_ok=True)
                else:
                    session.commit()
